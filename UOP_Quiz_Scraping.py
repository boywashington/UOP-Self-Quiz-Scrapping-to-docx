# This script is scraping quiz data from the UoPeople portal and exporting it to a Word document.
# Make sure to update the URL (see line 13)and COOKIES (see line 17) variables with your actual data before running.
# COOKIES information is copied from logged-in browser session (inspect element -> application -> cookies) and should be fresh to work properly.

import os
import re
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
import tkinter as tk
from tkinter import filedialog

# 1. Replace this with your actual URL
URL = "https://my.uopeople.edu/mod/quiz/review.php?attempt=22811690&cmid=512404"

# 2. MANDATORY: Replace with your actual active session cookie value
COOKIES = {
    'MoodleSession': 'ismus3ftpam2qec38ddtauekb7' 
}

HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
}

def format_choice_text(text):
    """
    Ensures there is a proper space between the choice letter (a, b, c, etc.) 
    and the text that follows it.
    """
    return re.sub(r'^([a-zA-Z]\.)\s*', r'\1 ', text)

def prompt_save_location():
    """Opens a native save dialog window to select the output .docx file path."""
    root = tk.Tk()
    root.withdraw()  # Hide the main small tkinter window
    root.attributes("-topmost", True)  # Bring the dialog to the front
    
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
        title="Select Save Location for Quiz Document",
        initialfile="Quiz_Scraped.docx"
    )
    return file_path

def scrape_quiz():
    print("Fetching quiz data from UoPeople portal...")
    try:
        response = requests.get(URL, cookies=COOKIES, headers=HEADERS, timeout=15)
    except Exception as e:
        print(f"Network error: {e}")
        return
    
    if response.status_code != 200:
        print(f"Failed to fetch page. Status code: {response.status_code}")
        print("Tip: Make sure your MoodleSession cookie is fresh and valid.")
        return

    soup = BeautifulSoup(response.text, 'html.parser')
    questions = soup.find_all('div', class_='que')
    
    if not questions:
        print("No questions found. Check your session cookie or URL.")
        return

    print(f"Successfully scraped {len(questions)} questions. Creating Word Document...")
    
    # Initialize Word Document
    doc = Document()
    
    # Styles Setup
    title = doc.add_heading('UoPeople Quiz Review Export', level=0)
    doc.add_paragraph(f"Source URL: {URL}\nQuestions Found: {len(questions)}")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    for index, q in enumerate(questions, 1):
        # 1. Add Question Heading
        q_heading = doc.add_heading(f"Question {index}", level=2)
        q_heading.paragraph_format.space_before = Pt(18)
        q_heading.paragraph_format.space_after = Pt(6)
        
        # 2. Add Question Text
        qtext_div = q.find('div', class_='qtext')
        if qtext_div:
            q_text = qtext_div.get_text(strip=True)
            p_qtext = doc.add_paragraph()
            run_qtext = p_qtext.add_run(q_text)
            run_qtext.bold = True
            p_qtext.paragraph_format.space_after = Pt(4) 
        else:
            doc.add_paragraph("[Could not parse question text]")
            
        # 3. Add Multiple Choices (Listed vertically, tight spacing)
        answer_block = q.find('div', class_=['ablock', 'answer'])
        if answer_block:
            choice_rows = answer_block.find_all('div', class_=['r0', 'r1'])
            if choice_rows:
                for row in choice_rows:
                    text_label = row.get_text(strip=True)
                    if text_label:
                        formatted_text = format_choice_text(text_label)
                        p_choice = doc.add_paragraph(formatted_text)
                        p_choice.paragraph_format.space_after = Pt(2) 
            else:
                # Fallback to labels
                labels = answer_block.find_all('label')
                for label in labels:
                    text_label = label.get_text(strip=True)
                    if text_label:
                        formatted_text = format_choice_text(text_label)
                        p_choice = doc.add_paragraph(formatted_text)
                        p_choice.paragraph_format.space_after = Pt(2) 
        else:
            doc.add_paragraph("[No multiple choice blocks found]")
                        
        # 4. Add Cleaned Feedback Content
        outcome_div = q.find('div', class_='outcome')
        if outcome_div:
            feedback = outcome_div.get_text(strip=True)
            if feedback:
                # Fixed: passing re.IGNORECASE properly as a parameter to avoid regex compilation issues
                cleaned_feedback = re.sub(r'^feedback\s*:?\s*', '', feedback, flags=re.IGNORECASE)
                
                if cleaned_feedback: 
                    p_feed = doc.add_paragraph()
                    p_feed.paragraph_format.space_before = Pt(8)
                    p_feed.paragraph_format.space_after = Pt(4)
                    run_feed = p_feed.add_run(cleaned_feedback)
                    run_feed.italic = True
            
        # Add a visual separator line between questions
        # doc.add_paragraph("_" * 60).paragraph_format.space_after = Pt(12)

    # Prompt user for save location
    print("\nOpening prompt window to select file save location...")
    save_path = prompt_save_location()
    
    if save_path:
        try:
            doc.save(save_path)
            print(f"\n[SUCCESS] Document saved successfully to:\n{os.path.abspath(save_path)}")
        except Exception as e:
            print(f"\n[ERROR] Failed to save file: {e}")
    else:
        print("\n[CANCELLED] File save operation cancelled by user.")

if __name__ == "__main__":
    scrape_quiz()